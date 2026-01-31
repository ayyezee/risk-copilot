"""Document generation service for creating DOCX output with formatting preservation.

This service handles:
1. DOCX input: Preserves original formatting while applying replacements
2. PDF input: Creates clean DOCX from extracted text
3. Track changes: Generates a changes report document
"""

import io
import re
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import BinaryIO

import structlog
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.services.document_ai_processor import TermReplacement

logger = structlog.get_logger()


@dataclass
class ReplacementMatch:
    """Tracks a single replacement occurrence in the document."""
    original_term: str
    replacement_term: str
    paragraph_index: int
    location_description: str
    reasoning: str
    confidence: float


@dataclass
class GenerationResult:
    """Result of document generation."""
    output_bytes: bytes
    output_filename: str
    content_type: str
    total_replacements_applied: int
    replacement_details: list[ReplacementMatch] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    source_format: str = "docx"  # "docx" or "pdf"


class DocumentGenerator:
    """Generates DOCX documents with applied replacements while preserving formatting."""

    # Common conjunctions that can create redundant phrases when terms are replaced
    DEDUP_CONJUNCTIONS = [
        r'\s+and\s+',
        r'\s+or\s+',
        r'\s+and/or\s+',
        r',\s+and\s+',
        r',\s+or\s+',
        r',\s+',
        r'\s*/\s*',
    ]

    def __init__(self) -> None:
        self.logger = structlog.get_logger()

    def apply_replacements_to_docx(
        self,
        input_file: BinaryIO,
        replacements: list[TermReplacement],
        case_sensitive: bool = False,
        highlight_changes: bool = True,
        min_confidence: float = 0.7,
    ) -> GenerationResult:
        """Apply term replacements to a DOCX file while preserving formatting.

        Args:
            input_file: File-like object containing the original DOCX
            replacements: List of term replacements to apply
            case_sensitive: Whether to match case-sensitively
            highlight_changes: Whether to highlight replaced text
            min_confidence: Minimum confidence threshold for applying replacements

        Returns:
            GenerationResult with the modified document
        """
        # Filter replacements by confidence
        applicable_replacements = [
            r for r in replacements if r.confidence >= min_confidence
        ]

        # Sort by original term length (longest first) to avoid partial replacements
        applicable_replacements.sort(key=lambda r: len(r.original_term), reverse=True)

        # Load the document
        doc = Document(input_file)

        replacement_matches: list[ReplacementMatch] = []
        total_replacements = 0

        # Process main document body paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            count, matches = self._apply_replacements_to_paragraph(
                paragraph=paragraph,
                replacements=applicable_replacements,
                case_sensitive=case_sensitive,
                highlight_changes=highlight_changes,
                para_index=para_idx,
                location="body",
            )
            total_replacements += count
            replacement_matches.extend(matches)

        # Process tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        count, matches = self._apply_replacements_to_paragraph(
                            paragraph=paragraph,
                            replacements=applicable_replacements,
                            case_sensitive=case_sensitive,
                            highlight_changes=highlight_changes,
                            para_index=para_idx,
                            location=f"table {table_idx + 1}, row {row_idx + 1}, cell {cell_idx + 1}",
                        )
                        total_replacements += count
                        replacement_matches.extend(matches)

        # Process headers
        for section in doc.sections:
            for header_type in ['header', 'first_page_header', 'even_page_header']:
                header = getattr(section, header_type, None)
                if header:
                    for para_idx, paragraph in enumerate(header.paragraphs):
                        count, matches = self._apply_replacements_to_paragraph(
                            paragraph=paragraph,
                            replacements=applicable_replacements,
                            case_sensitive=case_sensitive,
                            highlight_changes=highlight_changes,
                            para_index=para_idx,
                            location=f"{header_type}",
                        )
                        total_replacements += count
                        replacement_matches.extend(matches)

            # Process footers
            for footer_type in ['footer', 'first_page_footer', 'even_page_footer']:
                footer = getattr(section, footer_type, None)
                if footer:
                    for para_idx, paragraph in enumerate(footer.paragraphs):
                        count, matches = self._apply_replacements_to_paragraph(
                            paragraph=paragraph,
                            replacements=applicable_replacements,
                            case_sensitive=case_sensitive,
                            highlight_changes=highlight_changes,
                            para_index=para_idx,
                            location=f"{footer_type}",
                        )
                        total_replacements += count
                        replacement_matches.extend(matches)

        # Post-process: deduplicate redundant phrases
        # This handles cases like "Investment Subadvisor and Investment Subadvisor"
        replacement_targets = self._get_unique_replacement_targets(applicable_replacements)
        dedup_count = 0

        # Deduplicate main body paragraphs
        for paragraph in doc.paragraphs:
            dedup_count += self._deduplicate_paragraph(paragraph, replacement_targets)

        # Deduplicate tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        dedup_count += self._deduplicate_paragraph(paragraph, replacement_targets)

        # Deduplicate headers and footers
        for section in doc.sections:
            for header_type in ['header', 'first_page_header', 'even_page_header']:
                header = getattr(section, header_type, None)
                if header:
                    for paragraph in header.paragraphs:
                        dedup_count += self._deduplicate_paragraph(paragraph, replacement_targets)

            for footer_type in ['footer', 'first_page_footer', 'even_page_footer']:
                footer = getattr(section, footer_type, None)
                if footer:
                    for paragraph in footer.paragraphs:
                        dedup_count += self._deduplicate_paragraph(paragraph, replacement_targets)

        if dedup_count > 0:
            self.logger.info(
                "Deduplicated redundant phrases",
                dedup_count=dedup_count,
            )

        # Post-process: fix grammar issues (subject-verb agreement)
        # This handles cases like "The Investment Subadvisor are" -> "The Investment Subadvisor is"
        grammar_fix_count = 0

        for paragraph in doc.paragraphs:
            grammar_fix_count += self._fix_paragraph_grammar(paragraph, replacement_targets)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        grammar_fix_count += self._fix_paragraph_grammar(paragraph, replacement_targets)

        for section in doc.sections:
            for header_type in ['header', 'first_page_header', 'even_page_header']:
                header = getattr(section, header_type, None)
                if header:
                    for paragraph in header.paragraphs:
                        grammar_fix_count += self._fix_paragraph_grammar(paragraph, replacement_targets)

            for footer_type in ['footer', 'first_page_footer', 'even_page_footer']:
                footer = getattr(section, footer_type, None)
                if footer:
                    for paragraph in footer.paragraphs:
                        grammar_fix_count += self._fix_paragraph_grammar(paragraph, replacement_targets)

        if grammar_fix_count > 0:
            self.logger.info(
                "Fixed grammar issues after replacement",
                grammar_fix_count=grammar_fix_count,
            )

        # Post-process: final style pass (possessives, capitalization, remaining duplicates)
        style_fix_count = 0

        for paragraph in doc.paragraphs:
            style_fix_count += self._apply_final_style_to_paragraph(paragraph, replacement_targets)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        style_fix_count += self._apply_final_style_to_paragraph(paragraph, replacement_targets)

        for section in doc.sections:
            for header_type in ['header', 'first_page_header', 'even_page_header']:
                header = getattr(section, header_type, None)
                if header:
                    for paragraph in header.paragraphs:
                        style_fix_count += self._apply_final_style_to_paragraph(paragraph, replacement_targets)

            for footer_type in ['footer', 'first_page_footer', 'even_page_footer']:
                footer = getattr(section, footer_type, None)
                if footer:
                    for paragraph in footer.paragraphs:
                        style_fix_count += self._apply_final_style_to_paragraph(paragraph, replacement_targets)

        if style_fix_count > 0:
            self.logger.info(
                "Applied final style fixes",
                style_fix_count=style_fix_count,
            )

        # Save to bytes
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)

        self.logger.info(
            "DOCX generation complete",
            total_replacements=total_replacements,
            unique_terms_replaced=len(set(m.original_term for m in replacement_matches)),
        )

        return GenerationResult(
            output_bytes=output_buffer.getvalue(),
            output_filename=f"processed_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            total_replacements_applied=total_replacements,
            replacement_details=replacement_matches,
            source_format="docx",
        )

    def _apply_replacements_to_paragraph(
        self,
        paragraph: Paragraph,
        replacements: list[TermReplacement],
        case_sensitive: bool,
        highlight_changes: bool,
        para_index: int,
        location: str,
    ) -> tuple[int, list[ReplacementMatch]]:
        """Apply replacements to a single paragraph while preserving run formatting.

        This method handles the complex case where a term to replace might span
        multiple runs (e.g., "the <b>vendor</b>" where "vendor" is bold).

        Args:
            paragraph: The paragraph to process
            replacements: Replacements to apply
            case_sensitive: Whether to match case-sensitively
            highlight_changes: Whether to highlight changes
            para_index: Index of the paragraph
            location: Description of where this paragraph is

        Returns:
            Tuple of (replacement count, list of matches)
        """
        if not paragraph.runs:
            return 0, []

        # Get the full paragraph text
        full_text = paragraph.text
        if not full_text.strip():
            return 0, []

        total_count = 0
        matches: list[ReplacementMatch] = []

        # Process each replacement
        for replacement in replacements:
            flags = 0 if case_sensitive else re.IGNORECASE
            pattern = re.compile(
                r'\b' + re.escape(replacement.original_term) + r'\b',
                flags
            )

            # Check if the term exists in this paragraph
            if not pattern.search(full_text):
                continue

            # Apply the replacement using the run-aware method
            count = self._replace_in_paragraph_runs(
                paragraph=paragraph,
                pattern=pattern,
                replacement_text=replacement.replacement_term,
                highlight=highlight_changes,
            )

            if count > 0:
                total_count += count
                matches.append(ReplacementMatch(
                    original_term=replacement.original_term,
                    replacement_term=replacement.replacement_term,
                    paragraph_index=para_index,
                    location_description=location,
                    reasoning=replacement.reasoning,
                    confidence=replacement.confidence,
                ))

            # Update full_text for next iteration
            full_text = paragraph.text

        return total_count, matches

    def _replace_in_paragraph_runs(
        self,
        paragraph: Paragraph,
        pattern: re.Pattern,
        replacement_text: str,
        highlight: bool,
    ) -> int:
        """Replace text in paragraph runs while preserving formatting.

        This handles the case where the text to replace spans multiple runs.

        Args:
            paragraph: The paragraph to modify
            pattern: Compiled regex pattern for matching
            replacement_text: Text to replace matches with
            highlight: Whether to highlight the replacement

        Returns:
            Number of replacements made
        """
        # Build a map of character positions to runs
        runs = paragraph.runs
        if not runs:
            return 0

        # Concatenate all run texts and track boundaries
        full_text = ""
        run_boundaries = []  # List of (start, end, run_index)

        for i, run in enumerate(runs):
            start = len(full_text)
            full_text += run.text
            end = len(full_text)
            run_boundaries.append((start, end, i))

        # Find all matches
        matches = list(pattern.finditer(full_text))
        if not matches:
            return 0

        # Process matches in reverse order to preserve positions
        for match in reversed(matches):
            match_start = match.start()
            match_end = match.end()

            # Find which runs this match spans
            affected_runs = []
            for start, end, run_idx in run_boundaries:
                if start < match_end and end > match_start:
                    # Calculate the portion of this run that's affected
                    local_start = max(0, match_start - start)
                    local_end = min(end - start, match_end - start)
                    affected_runs.append((run_idx, local_start, local_end))

            if not affected_runs:
                continue

            # Handle single-run case (most common)
            if len(affected_runs) == 1:
                run_idx, local_start, local_end = affected_runs[0]
                run = runs[run_idx]
                old_text = run.text
                new_text = old_text[:local_start] + replacement_text + old_text[local_end:]
                run.text = new_text

                if highlight:
                    self._highlight_run(run)

            else:
                # Multi-run case: put replacement in first run, clear others
                first_run_idx, first_local_start, first_local_end = affected_runs[0]
                first_run = runs[first_run_idx]

                # Preserve formatting from first run
                old_text = first_run.text
                first_run.text = old_text[:first_local_start] + replacement_text

                if highlight:
                    self._highlight_run(first_run)

                # Clear the matched portions from subsequent runs
                for run_idx, local_start, local_end in affected_runs[1:]:
                    run = runs[run_idx]
                    old_text = run.text
                    # Keep text before and after the match portion
                    run.text = old_text[:local_start] + old_text[local_end:]

                # Handle text after the match in the last affected run
                last_run_idx, last_local_start, last_local_end = affected_runs[-1]
                if last_run_idx != first_run_idx:
                    last_run = runs[last_run_idx]
                    # Text after match is already preserved from the clear step above

        # Recalculate boundaries and return count
        return len(matches)

    def _highlight_run(self, run: Run) -> None:
        """Apply yellow highlight to a run to mark it as changed."""
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    def _get_unique_replacement_targets(
        self,
        replacements: list[TermReplacement],
    ) -> set[str]:
        """Extract unique replacement target terms."""
        return {r.replacement_term for r in replacements}

    # Grammar fixes for subject-verb agreement when replacing plural subjects with singular
    # Format: (pattern, replacement) - pattern matches the incorrect form after replacement
    # NOTE: These patterns require the entity to be THE SUBJECT (start of sentence or after punctuation)
    # They should not match when the entity is inside a prepositional phrase (e.g., "made by the XYZ Series")
    GRAMMAR_FIXES = [
        # "The [Singular] are" -> "The [Singular] is"
        # Only match when "The" is at word boundary (start of sentence/clause)
        (r'\b(The\s+Investment\s+Subadvisor)\s+are\b', r'\1 is'),
        (r'\b(The\s+Investment\s+Subadvisor)\s+were\b', r'\1 was'),
        (r'\b(The\s+Investment\s+Subadvisor)\s+have\b', r'\1 has'),
        (r'\b(The\s+Investment\s+Subadvisor)\s+do\b', r'\1 does'),
        # Same for XYZ Series with "The" prefix (indicates it's the subject)
        (r'\b(The\s+XYZ\s+Series)\s+are\b', r'\1 is'),
        (r'\b(The\s+XYZ\s+Series)\s+were\b', r'\1 was'),
        (r'\b(The\s+XYZ\s+Series)\s+have\b', r'\1 has'),
        # NOTE: Removed patterns for "XYZ Series are" without "The" because these
        # can incorrectly match when the true subject is plural (e.g., "Investments made by XYZ Series are")
        # "their" -> "its" for singular entities
        (r'\b(The\s+Investment\s+Subadvisor)\s+and\s+their\b', r'\1 and its'),
        (r'\b(Investment\s+Subadvisor)\s+in\s+their\b', r'\1 in its'),
        (r'\b(XYZ\s+Series)\s+and\s+their\b', r'\1 and its'),
    ]

    def _deduplicate_redundant_text(
        self,
        text: str,
        replacement_targets: set[str],
    ) -> tuple[str, int]:
        """Remove redundant phrases from text where the same term appears multiple times.

        For example: "Investment Subadvisor and Investment Subadvisor" -> "Investment Subadvisor"

        Args:
            text: Text to process
            replacement_targets: Set of replacement terms that might appear redundantly

        Returns:
            Tuple of (processed text, number of deduplication changes made)
        """
        modified_text = text
        total_changes = 0

        for target in replacement_targets:
            # Skip short terms to avoid false positives
            if len(target) < 3:
                continue

            escaped_target = re.escape(target)

            # Pattern 1: Basic "Target [conjunction] Target" -> "Target"
            for conjunction in self.DEDUP_CONJUNCTIONS:
                pattern = re.compile(
                    rf'\b({escaped_target})({conjunction})({escaped_target})\b',
                    re.IGNORECASE
                )

                while pattern.search(modified_text):
                    modified_text = pattern.sub(r'\1', modified_text)
                    total_changes += 1

            # Pattern 2: "the Target and the Target" -> "the Target"
            pattern_with_article = re.compile(
                rf'\b(the\s+{escaped_target})(\s+and\s+|\s+or\s+|\s+and/or\s+)(the\s+{escaped_target})\b',
                re.IGNORECASE
            )

            while pattern_with_article.search(modified_text):
                modified_text = pattern_with_article.sub(r'\1', modified_text)
                total_changes += 1

            # Pattern 3: "The Target, the Target," -> "The Target," (with articles and commas)
            # Handles: "The Investment Subadvisor, the Investment Subadvisor, and"
            pattern_comma_article = re.compile(
                rf'\b(The\s+{escaped_target}),\s+the\s+{escaped_target},',
                re.IGNORECASE
            )

            while pattern_comma_article.search(modified_text):
                modified_text = pattern_comma_article.sub(r'\1,', modified_text)
                total_changes += 1

            # Pattern 4: "the Target, the Target" (lowercase, comma separated)
            pattern_comma_lower = re.compile(
                rf'\b(the\s+{escaped_target}),\s+the\s+{escaped_target}\b',
                re.IGNORECASE
            )

            while pattern_comma_lower.search(modified_text):
                modified_text = pattern_comma_lower.sub(r'\1', modified_text)
                total_changes += 1

            # Pattern 5: "Target, Target," or "Target, Target and" -> "Target,"/"Target and"
            pattern_bare_comma = re.compile(
                rf'\b({escaped_target}),\s+{escaped_target}(,|\s+and\b)',
                re.IGNORECASE
            )

            while pattern_bare_comma.search(modified_text):
                modified_text = pattern_bare_comma.sub(r'\1\2', modified_text)
                total_changes += 1

            # Pattern 6: Handle repeated words like "Series Series" -> "Series"
            repeated_word_pattern = re.compile(
                rf'\b({escaped_target})\s+\1\b',
                re.IGNORECASE
            )
            while repeated_word_pattern.search(modified_text):
                modified_text = repeated_word_pattern.sub(r'\1', modified_text)
                total_changes += 1

            # Pattern 7: "The Target, the Target and" -> "The Target and"
            # (article + comma + article + "and")
            pattern_comma_and = re.compile(
                rf'\b(The\s+{escaped_target}),\s+the\s+{escaped_target}\s+and\b',
                re.IGNORECASE
            )

            while pattern_comma_and.search(modified_text):
                modified_text = pattern_comma_and.sub(r'\1 and', modified_text)
                total_changes += 1

            # Pattern 8: Three in a row "Target, Target, and Target" -> "Target"
            pattern_three = re.compile(
                rf'\b({escaped_target}),\s+{escaped_target},\s+and\s+{escaped_target}\b',
                re.IGNORECASE
            )

            while pattern_three.search(modified_text):
                modified_text = pattern_three.sub(r'\1', modified_text)
                total_changes += 1

        return modified_text, total_changes

    def _fix_grammar_after_replacement(
        self,
        text: str,
        replacement_targets: set[str],
    ) -> tuple[str, int]:
        """Fix grammatical errors that occur when plural terms are replaced with singular.

        For example:
        - "The General Partner and Investment Manager are" becomes
          "The Investment Subadvisor are" after deduplication
        - This should be "The Investment Subadvisor is"

        Also fixes subject-verb agreement where the true subject is not the replacement term:
        - "Investments made by the XYZ Series is" -> "Investments made by the XYZ Series are"

        Args:
            text: Text after replacements and deduplication
            replacement_targets: Set of replacement terms

        Returns:
            Tuple of (corrected text, number of grammar fixes made)
        """
        modified_text = text
        total_fixes = 0

        # FIRST: Apply static grammar fixes for singular subjects
        # These fix cases like "The XYZ Series are" -> "The XYZ Series is"
        for pattern_str, replacement in self.GRAMMAR_FIXES:
            pattern = re.compile(pattern_str, re.IGNORECASE)
            matches = pattern.findall(modified_text)
            if matches:
                modified_text = pattern.sub(replacement, modified_text)
                total_fixes += len(matches)

        # Dynamic grammar fixes based on replacement targets
        for target in replacement_targets:
            if len(target) < 3:
                continue

            escaped_target = re.escape(target)

            # Fix "The [Target] are" -> "The [Target] is" (for singular targets)
            # This applies to terms that look singular (don't end in 's' or end in specific patterns)
            is_likely_singular = (
                not target.lower().endswith('s') or
                target.lower().endswith('series') or
                target.lower().endswith('subadvisor') or
                'limited partners' not in target.lower()
            )

            if is_likely_singular:
                # Fix subject-verb agreement
                verb_fixes = [
                    (rf'\b(The\s+{escaped_target})\s+are\b', r'\1 is'),
                    (rf'\b(The\s+{escaped_target})\s+were\b', r'\1 was'),
                    (rf'\b(The\s+{escaped_target})\s+have\b', r'\1 has'),
                    (rf'\b(The\s+{escaped_target})\s+do\b', r'\1 does'),
                    (rf'\b({escaped_target})\s+are\s+under\b', r'\1 is under'),
                    (rf'\b({escaped_target})\s+are\s+not\b', r'\1 is not'),
                    (rf'\b({escaped_target})\s+are\s+responsible\b', r'\1 is responsible'),
                    (rf'\b({escaped_target})\s+are\s+entitled\b', r'\1 is entitled'),
                    (rf'\b({escaped_target})\s+are\s+required\b', r'\1 is required'),
                    (rf'\b({escaped_target})\s+are\s+authorized\b', r'\1 is authorized'),
                    (rf'\b({escaped_target})\s+are\s+obligated\b', r'\1 is obligated'),
                    (rf'\b({escaped_target})\s+are\s+permitted\b', r'\1 is permitted'),
                    (rf'\b({escaped_target})\s+have\s+no\b', r'\1 has no'),
                    (rf'\b({escaped_target})\s+have\s+the\b', r'\1 has the'),
                    (rf'\b({escaped_target})\s+have\s+full\b', r'\1 has full'),
                    (rf'\b({escaped_target})\s+have\s+sole\b', r'\1 has sole'),
                    (rf'\b({escaped_target})\s+do\s+not\b', r'\1 does not'),
                ]

                for pattern_str, replacement in verb_fixes:
                    pattern = re.compile(pattern_str, re.IGNORECASE)
                    if pattern.search(modified_text):
                        modified_text = pattern.sub(replacement, modified_text)
                        total_fixes += 1

                # Fix possessive pronouns: "their" -> "its" when referring to singular entity
                pronoun_fixes = [
                    # Specific phrases
                    (rf'({escaped_target})\s+in\s+their\s+sole\b', r'\1 in its sole'),
                    (rf'({escaped_target})\s+at\s+their\s+discretion\b', r'\1 at its discretion'),
                    (rf'({escaped_target})\s+and\s+their\s+affiliates\b', r'\1 and its affiliates'),
                    # General cases - "their" followed by common nouns
                    (rf'({escaped_target})\s+(\w+\s+)*their\s+(investment|capital|contributions?|interests?|objectives?|performance)\b',
                     rf'\1 \2its \3'),
                ]

                for pattern_str, replacement in pronoun_fixes:
                    pattern = re.compile(pattern_str, re.IGNORECASE)
                    if pattern.search(modified_text):
                        modified_text = pattern.sub(replacement, modified_text)
                        total_fixes += 1

        # LAST: Fix subject-verb agreement where a PLURAL subject is followed by
        # a prepositional phrase containing a singular entity, then incorrect singular verb
        # This must come AFTER the singular subject fixes above to correct any over-corrections
        # Pattern: "Investments [made by X] is" -> "Investments [made by X] are"
        plural_subject_fixes = [
            # "Investments made by the XYZ Series is expected" -> "are expected"
            (r'\b(Investments\s+made\s+by\s+(?:the\s+)?[A-Z][A-Za-z\s]+?)\s+is\s+(expected|subject|primarily)\b', r'\1 are \2'),
            # More general: "Investments in X is" -> "are"
            (r'\b(Investments\s+in\s+(?:the\s+)?[A-Z][A-Za-z\s]+?)\s+is\b', r'\1 are'),
            (r'\b(Securities\s+(?:held\s+by|in)\s+[^,]+?)\s+is\b', r'\1 are'),
            (r'\b(Assets\s+(?:held\s+by|in)\s+[^,]+?)\s+is\b', r'\1 are'),
        ]

        for pattern_str, replacement in plural_subject_fixes:
            pattern = re.compile(pattern_str, re.IGNORECASE)
            if pattern.search(modified_text):
                modified_text = pattern.sub(replacement, modified_text)
                total_fixes += 1

        return modified_text, total_fixes

    def _final_style_pass(
        self,
        text: str,
        replacement_targets: set[str],
    ) -> tuple[str, int]:
        """Final pass to fix style issues after all other processing.

        This handles:
        1. Possessive form for words ending in 's' (Series's -> Series')
        2. Pronoun agreement: "their" -> "its" for singular entities
        3. Sentence capitalization: lowercase "the" at sentence start -> "The"
        4. Catch any remaining repeated words (including "Series Series")
        5. Fix duplicate entities in conjunctions ("X nor X" -> "X")

        Args:
            text: Text after all other processing
            replacement_targets: Set of replacement terms

        Returns:
            Tuple of (corrected text, number of style fixes made)
        """
        modified_text = text
        total_fixes = 0

        # Normalize targets - strip "the " prefix if present for pattern matching
        normalized_targets = set()
        for target in replacement_targets:
            normalized_targets.add(target)
            # Also add version without "the " prefix
            if target.lower().startswith('the '):
                normalized_targets.add(target[4:])

        # Fix 0: Handle duplicate entities in "neither...nor" and similar constructions
        # "Neither the Investment Subadvisor nor the Investment Subadvisor" -> "The Investment Subadvisor"
        for target in normalized_targets:
            escaped_target = re.escape(target)

            # Pattern: "Neither the X nor the X" -> "The X"
            neither_nor_pattern = re.compile(
                rf'\bNeither\s+(the\s+)?{escaped_target}\s+nor\s+(the\s+)?{escaped_target}\b',
                re.IGNORECASE
            )
            if neither_nor_pattern.search(modified_text):
                # Replace with just "The X" (capitalized since it starts the clause)
                modified_text = neither_nor_pattern.sub(f'The {target}' if not target.lower().startswith('the ') else target.title() if target[0].islower() else target, modified_text)
                total_fixes += 1

            # Pattern: "either X or X" -> "X"
            either_or_pattern = re.compile(
                rf'\beither\s+(the\s+)?{escaped_target}\s+or\s+(the\s+)?{escaped_target}\b',
                re.IGNORECASE
            )
            if either_or_pattern.search(modified_text):
                modified_text = either_or_pattern.sub(target, modified_text)
                total_fixes += 1

            # Pattern: ", and their respective" when "their" refers to singular -> ", and its"
            respective_pattern = re.compile(
                rf'\b({escaped_target}),?\s+and\s+their\s+respective\b',
                re.IGNORECASE
            )
            if respective_pattern.search(modified_text):
                modified_text = respective_pattern.sub(r'\1 and its', modified_text)
                total_fixes += 1

        # Fix 1: Possessive form for words ending in 's'
        # "XYZ Series's" -> "XYZ Series'" (just apostrophe, not apostrophe-s)
        # Match any word ending in 's' or 'S' followed by 's
        # Also handle multi-word phrases ending in 's': "XYZ Series's" -> "XYZ Series'"
        possessive_s_pattern = re.compile(r"(\w+[sS])'s\b")

        def fix_possessive(match):
            word = match.group(1)
            return f"{word}'"

        new_text = possessive_s_pattern.sub(fix_possessive, modified_text)
        if new_text != modified_text:
            total_fixes += 1
            modified_text = new_text

        # Fix 2: Pronoun agreement - "their" -> "its" for singular entities
        # This applies when a singular entity (XYZ Series, Investment Subadvisor) is followed by "their"
        for target in normalized_targets:
            escaped_target = re.escape(target)

            # Skip targets that are genuinely plural (e.g., "Limited Partners")
            if 'partners' in target.lower() and 'series' not in target.lower():
                continue

            # Pattern: [the] Target ... their (within reasonable distance)
            # We look for "their" that refers back to the singular target
            their_patterns = [
                # Direct: "the XYZ Series may lose... their"
                (rf'\b(the\s+{escaped_target})\s+(may|will|should|could|would|can|cannot|must)([^.]*?)\s+their\b', r'\1 \2\3 its'),
                # Direct: "XYZ Series may lose... their"
                (rf'\b({escaped_target})\s+(may|will|should|could|would|can|cannot|must)([^.]*?)\s+their\b', r'\1 \2\3 its'),
                # "the XYZ Series should be prepared to bear... their investment"
                (rf'\b(the\s+{escaped_target})([^.]*?)\s+their\s+(investment|interests?|capital|contributions?)\b', r'\1\2 its \3'),
                # Without "the": "XYZ Series... their investment"
                (rf'\b({escaped_target})([^.]*?)\s+their\s+(investment|interests?|capital|contributions?)\b', r'\1\2 its \3'),
            ]

            for pattern_str, replacement in their_patterns:
                pattern = re.compile(pattern_str, re.IGNORECASE)
                if pattern.search(modified_text):
                    modified_text = pattern.sub(replacement, modified_text)
                    total_fixes += 1

        # Fix 3: Sentence capitalization - lowercase "the" at sentence start
        # Pattern: sentence boundary (. ! ? :) followed by space and lowercase "the"
        sentence_start_pattern = re.compile(r'([.!?:]\s+)(the\s+)', re.MULTILINE)

        def capitalize_sentence_start(match):
            punctuation = match.group(1)
            the_word = match.group(2)
            return punctuation + 'T' + the_word[1:]

        new_text = sentence_start_pattern.sub(capitalize_sentence_start, modified_text)
        if new_text != modified_text:
            total_fixes += 1
            modified_text = new_text

        # Fix 4: Paragraph/text start - "the" at very beginning should be "The"
        if modified_text.startswith('the '):
            modified_text = 'The ' + modified_text[4:]
            total_fixes += 1

        # Fix 5: Catch any remaining repeated words (e.g., "Series Series")
        # This handles cases like "Series Series Limited Partners"
        repeated_word_pattern = re.compile(r'\b(\w+)\s+\1\b', re.IGNORECASE)

        while repeated_word_pattern.search(modified_text):
            modified_text = repeated_word_pattern.sub(r'\1', modified_text)
            total_fixes += 1

        return modified_text, total_fixes

    def _apply_final_style_to_paragraph(
        self,
        paragraph: Paragraph,
        replacement_targets: set[str],
    ) -> int:
        """Apply final style fixes to a paragraph's runs.

        For multi-run paragraphs, we consolidate all text, apply fixes,
        then redistribute back to preserve formatting boundaries.

        Args:
            paragraph: The paragraph to process
            replacement_targets: Set of replacement terms

        Returns:
            Number of style fixes made
        """
        if not paragraph.runs:
            return 0

        full_text = paragraph.text
        if not full_text.strip():
            return 0

        # Check if style fixes are needed on the full text
        fixed_text, fixes = self._final_style_pass(full_text, replacement_targets)

        if fixes == 0:
            return 0

        # Apply the fixes to runs
        if len(paragraph.runs) == 1:
            paragraph.runs[0].text = fixed_text
        else:
            # For multiple runs, we need to map the fixed text back to runs
            # Strategy: Track character positions and redistribute
            runs = paragraph.runs
            original_lengths = [len(run.text) for run in runs]
            total_original = sum(original_lengths)
            total_fixed = len(fixed_text)

            # If lengths are similar, try to preserve run boundaries
            if abs(total_fixed - total_original) < 50:
                # Calculate position mapping
                pos = 0
                for i, run in enumerate(runs):
                    run_len = original_lengths[i]
                    if pos + run_len <= len(fixed_text):
                        run.text = fixed_text[pos:pos + run_len]
                    elif pos < len(fixed_text):
                        run.text = fixed_text[pos:]
                    else:
                        run.text = ""
                    pos += run_len

                # Handle any remaining text (put in last non-empty run)
                if pos < len(fixed_text):
                    for run in reversed(runs):
                        if run.text:
                            run.text += fixed_text[pos:]
                            break
            else:
                # Significant length change - put all text in first run, clear others
                runs[0].text = fixed_text
                for run in runs[1:]:
                    run.text = ""

        return fixes

    def _deduplicate_paragraph(
        self,
        paragraph: Paragraph,
        replacement_targets: set[str],
    ) -> int:
        """Apply deduplication to a paragraph's runs.

        For multi-run paragraphs, we consolidate all text, apply fixes,
        then redistribute back to preserve formatting boundaries.

        Args:
            paragraph: The paragraph to process
            replacement_targets: Set of replacement terms that might appear redundantly

        Returns:
            Number of deduplication changes made
        """
        if not paragraph.runs:
            return 0

        # Get the full paragraph text
        full_text = paragraph.text
        if not full_text.strip():
            return 0

        # Check if deduplication is needed
        deduped_text, changes = self._deduplicate_redundant_text(full_text, replacement_targets)

        if changes == 0:
            return 0

        # Apply the fixes to runs
        if len(paragraph.runs) == 1:
            paragraph.runs[0].text = deduped_text
        else:
            # For multiple runs, consolidate and redistribute
            runs = paragraph.runs
            original_lengths = [len(run.text) for run in runs]
            total_original = sum(original_lengths)
            total_fixed = len(deduped_text)

            if abs(total_fixed - total_original) < 50:
                pos = 0
                for i, run in enumerate(runs):
                    run_len = original_lengths[i]
                    if pos + run_len <= len(deduped_text):
                        run.text = deduped_text[pos:pos + run_len]
                    elif pos < len(deduped_text):
                        run.text = deduped_text[pos:]
                    else:
                        run.text = ""
                    pos += run_len

                if pos < len(deduped_text):
                    for run in reversed(runs):
                        if run.text:
                            run.text += deduped_text[pos:]
                            break
            else:
                runs[0].text = deduped_text
                for run in runs[1:]:
                    run.text = ""

        return changes

    def _fix_paragraph_grammar(
        self,
        paragraph: Paragraph,
        replacement_targets: set[str],
    ) -> int:
        """Apply grammar fixes to a paragraph's runs.

        For multi-run paragraphs, we consolidate all text, apply fixes,
        then redistribute back to preserve formatting boundaries.

        Args:
            paragraph: The paragraph to process
            replacement_targets: Set of replacement terms

        Returns:
            Number of grammar fixes made
        """
        if not paragraph.runs:
            return 0

        full_text = paragraph.text
        if not full_text.strip():
            return 0

        # Check if grammar fixes are needed
        fixed_text, fixes = self._fix_grammar_after_replacement(full_text, replacement_targets)

        if fixes == 0:
            return 0

        # Apply the fixes to runs
        if len(paragraph.runs) == 1:
            paragraph.runs[0].text = fixed_text
        else:
            # For multiple runs, consolidate and redistribute
            runs = paragraph.runs
            original_lengths = [len(run.text) for run in runs]
            total_original = sum(original_lengths)
            total_fixed = len(fixed_text)

            if abs(total_fixed - total_original) < 50:
                pos = 0
                for i, run in enumerate(runs):
                    run_len = original_lengths[i]
                    if pos + run_len <= len(fixed_text):
                        run.text = fixed_text[pos:pos + run_len]
                    elif pos < len(fixed_text):
                        run.text = fixed_text[pos:]
                    else:
                        run.text = ""
                    pos += run_len

                if pos < len(fixed_text):
                    for run in reversed(runs):
                        if run.text:
                            run.text += fixed_text[pos:]
                            break
            else:
                runs[0].text = fixed_text
                for run in runs[1:]:
                    run.text = ""

        return fixes

    def create_docx_from_text(
        self,
        text: str,
        replacements: list[TermReplacement],
        original_filename: str,
        case_sensitive: bool = False,
        highlight_changes: bool = True,
        min_confidence: float = 0.7,
    ) -> GenerationResult:
        """Create a DOCX document from plain text (e.g., extracted from PDF).

        This is used when the original format can't preserve formatting (PDF).

        Args:
            text: The extracted text content
            replacements: List of term replacements to apply
            original_filename: Name of the original file
            case_sensitive: Whether to match case-sensitively
            highlight_changes: Whether to highlight replaced text
            min_confidence: Minimum confidence threshold

        Returns:
            GenerationResult with the new document
        """
        # Filter replacements by confidence
        applicable_replacements = [
            r for r in replacements if r.confidence >= min_confidence
        ]
        applicable_replacements.sort(key=lambda r: len(r.original_term), reverse=True)

        # Create a new document
        doc = Document()

        # Add a header noting this is a conversion
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(
            f"Converted from: {original_filename}\n"
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"Note: Original formatting could not be preserved."
        )
        header_run.font.size = Pt(9)
        header_run.font.italic = True
        header_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # Spacer

        # Apply replacements to the text
        modified_text = text
        replacement_matches: list[ReplacementMatch] = []
        total_replacements = 0

        for replacement in applicable_replacements:
            flags = 0 if case_sensitive else re.IGNORECASE
            pattern = re.compile(
                r'\b' + re.escape(replacement.original_term) + r'\b',
                flags
            )

            matches = list(pattern.finditer(modified_text))
            if matches:
                modified_text = pattern.sub(replacement.replacement_term, modified_text)
                total_replacements += len(matches)
                replacement_matches.append(ReplacementMatch(
                    original_term=replacement.original_term,
                    replacement_term=replacement.replacement_term,
                    paragraph_index=0,
                    location_description="document body",
                    reasoning=replacement.reasoning,
                    confidence=replacement.confidence,
                ))

        # Post-process: deduplicate redundant phrases
        replacement_targets = self._get_unique_replacement_targets(applicable_replacements)
        modified_text, dedup_count = self._deduplicate_redundant_text(modified_text, replacement_targets)

        if dedup_count > 0:
            self.logger.info(
                "Deduplicated redundant phrases in text",
                dedup_count=dedup_count,
            )

        # Post-process: fix grammar issues (subject-verb agreement)
        modified_text, grammar_fix_count = self._fix_grammar_after_replacement(
            modified_text, replacement_targets
        )

        if grammar_fix_count > 0:
            self.logger.info(
                "Fixed grammar issues in text",
                grammar_fix_count=grammar_fix_count,
            )

        # Post-process: final style pass (possessives, capitalization, remaining duplicates)
        modified_text, style_fix_count = self._final_style_pass(
            modified_text, replacement_targets
        )

        if style_fix_count > 0:
            self.logger.info(
                "Applied final style fixes to text",
                style_fix_count=style_fix_count,
            )

        # Add the text as paragraphs (split on double newlines)
        paragraphs = modified_text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Handle single newlines within the paragraph
                lines = para_text.split('\n')
                para = doc.add_paragraph()
                for i, line in enumerate(lines):
                    if i > 0:
                        para.add_run('\n')
                    run = para.add_run(line)

                    # If highlighting, we need to check if this line contains replacements
                    if highlight_changes:
                        for rep in applicable_replacements:
                            if rep.replacement_term.lower() in line.lower():
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                break

        # Save to bytes
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)

        warnings = [
            "Original document was PDF - formatting could not be preserved.",
            "Document has been converted to a clean DOCX format.",
        ]

        self.logger.info(
            "PDF-to-DOCX conversion complete",
            total_replacements=total_replacements,
            original_file=original_filename,
        )

        return GenerationResult(
            output_bytes=output_buffer.getvalue(),
            output_filename=f"converted_{Path(original_filename).stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            total_replacements_applied=total_replacements,
            replacement_details=replacement_matches,
            warnings=warnings,
            source_format="pdf",
        )

    def generate_changes_report(
        self,
        replacement_matches: list[ReplacementMatch],
        original_filename: str,
        document_summary: str | None = None,
    ) -> GenerationResult:
        """Generate a separate document detailing all changes made.

        This serves as an audit trail and review document.

        Args:
            replacement_matches: All replacements that were applied
            original_filename: Name of the original file
            document_summary: Optional AI-generated summary

        Returns:
            GenerationResult with the changes report document
        """
        doc = Document()

        # Title
        title = doc.add_heading("Document Changes Report", level=0)

        # Metadata section
        doc.add_heading("Document Information", level=1)
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'

        cells = info_table.rows[0].cells
        cells[0].text = "Original File"
        cells[1].text = original_filename

        cells = info_table.rows[1].cells
        cells[0].text = "Processed Date"
        cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        cells = info_table.rows[2].cells
        cells[0].text = "Total Replacements"
        cells[1].text = str(len(replacement_matches))

        doc.add_paragraph()  # Spacer

        # Summary section
        if document_summary:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(document_summary)

        # Changes detail section
        doc.add_heading("Detailed Changes", level=1)

        if not replacement_matches:
            doc.add_paragraph("No replacements were applied to this document.")
        else:
            # Group by original term
            from collections import defaultdict
            grouped = defaultdict(list)
            for match in replacement_matches:
                grouped[match.original_term].append(match)

            for original_term, matches in grouped.items():
                # Term header
                term_para = doc.add_paragraph()
                term_run = term_para.add_run(f"'{original_term}' → '{matches[0].replacement_term}'")
                term_run.bold = True
                term_run.font.size = Pt(12)

                # Reasoning
                reason_para = doc.add_paragraph()
                reason_para.add_run("Reasoning: ").bold = True
                reason_para.add_run(matches[0].reasoning)

                # Confidence
                conf_para = doc.add_paragraph()
                conf_para.add_run("Confidence: ").bold = True
                confidence_pct = f"{matches[0].confidence * 100:.0f}%"
                conf_para.add_run(confidence_pct)

                # Locations
                locations = list(set(m.location_description for m in matches))
                if locations:
                    loc_para = doc.add_paragraph()
                    loc_para.add_run("Locations: ").bold = True
                    loc_para.add_run(", ".join(locations))

                doc.add_paragraph()  # Spacer between terms

        # Save to bytes
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)

        return GenerationResult(
            output_bytes=output_buffer.getvalue(),
            output_filename=f"changes_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            total_replacements_applied=len(replacement_matches),
            replacement_details=replacement_matches,
            source_format="report",
        )


# Singleton instance
_document_generator_instance: DocumentGenerator | None = None


def get_document_generator() -> DocumentGenerator:
    """Get document generator singleton instance."""
    global _document_generator_instance
    if _document_generator_instance is None:
        _document_generator_instance = DocumentGenerator()
    return _document_generator_instance
